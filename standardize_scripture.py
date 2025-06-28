#!/usr/bin/env python3
"""
standardize_scripture.py - A script to standardize Bible references in Word documents.

This script processes Word documents to standardize Bible references to the format "Book Chapter:Verse" (e.g., "John 3:16").
It handles various abbreviations, creates backups before making changes, and includes proper error handling.
Verse ranges use hyphens (e.g., "John 3:16-18") and multiple verses use commas (e.g., "John 3:16,18").

Dependencies:
    - python-docx: pip install python-docx
    - re: Standard library for regular expressions
    - os, sys, shutil: Standard libraries for file operations
    - datetime: Standard library for timestamps
    - argparse: Standard library for command-line arguments
"""

import os
import re
import sys
import shutil
import argparse
from datetime import datetime
from typing import Dict, List, Tuple, Pattern, Optional
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run

# Dictionary mapping Bible book abbreviations to standardized names
BIBLE_BOOKS = {
    # Old Testament
    'gen': 'Genesis', 'ge': 'Genesis', 'gn': 'Genesis', 'genesis': 'Genesis',
    'exod': 'Exodus', 'ex': 'Exodus', 'exo': 'Exodus', 'exodus': 'Exodus',
    'lev': 'Leviticus', 'le': 'Leviticus', 'lv': 'Leviticus', 'leviticus': 'Leviticus',
    'num': 'Numbers', 'nu': 'Numbers', 'nm': 'Numbers', 'numbers': 'Numbers',
    'deut': 'Deuteronomy', 'dt': 'Deuteronomy', 'de': 'Deuteronomy', 'deuteronomy': 'Deuteronomy',
    'josh': 'Joshua', 'jos': 'Joshua', 'jsh': 'Joshua', 'joshua': 'Joshua',
    'judg': 'Judges', 'jdg': 'Judges', 'jg': 'Judges', 'judges': 'Judges',
    'ruth': 'Ruth', 'rth': 'Ruth', 'ru': 'Ruth',
    '1 sam': '1 Samuel', '1sam': '1 Samuel', '1 sa': '1 Samuel', '1sa': '1 Samuel', '1 samuel': '1 Samuel',
    '2 sam': '2 Samuel', '2sam': '2 Samuel', '2 sa': '2 Samuel', '2sa': '2 Samuel', '2 samuel': '2 Samuel',
    '1 kgs': '1 Kings', '1kgs': '1 Kings', '1 ki': '1 Kings', '1ki': '1 Kings', '1 kings': '1 Kings',
    '2 kgs': '2 Kings', '2kgs': '2 Kings', '2 ki': '2 Kings', '2ki': '2 Kings', '2 kings': '2 Kings',
    '1 chr': '1 Chronicles', '1chr': '1 Chronicles', '1 ch': '1 Chronicles', '1ch': '1 Chronicles', '1 chronicles': '1 Chronicles',
    '2 chr': '2 Chronicles', '2chr': '2 Chronicles', '2 ch': '2 Chronicles', '2ch': '2 Chronicles', '2 chronicles': '2 Chronicles',
    'ezra': 'Ezra', 'ezr': 'Ezra', 'ez': 'Ezra',
    'neh': 'Nehemiah', 'ne': 'Nehemiah', 'nehemiah': 'Nehemiah',
    'esth': 'Esther', 'est': 'Esther', 'es': 'Esther', 'esther': 'Esther',
    'job': 'Job', 'jb': 'Job',
    'ps': 'Psalm', 'psa': 'Psalm', 'psm': 'Psalm', 'psalm': 'Psalm', 'ps.': 'Psalm',
    'pss': 'Psalms', 'psalms': 'Psalms',
    'prov': 'Proverbs', 'pro': 'Proverbs', 'pr': 'Proverbs', 'prv': 'Proverbs', 'proverbs': 'Proverbs',
    'eccl': 'Ecclesiastes', 'ecc': 'Ecclesiastes', 'ec': 'Ecclesiastes', 'ecclesiastes': 'Ecclesiastes',
    'song': 'Song of Solomon', 'sos': 'Song of Solomon', 'ss': 'Song of Solomon', 'song of solomon': 'Song of Solomon',
    'isa': 'Isaiah', 'is': 'Isaiah', 'isaiah': 'Isaiah',
    'jer': 'Jeremiah', 'je': 'Jeremiah', 'jeremiah': 'Jeremiah',
    'lam': 'Lamentations', 'la': 'Lamentations', 'lamentations': 'Lamentations',
    'ezek': 'Ezekiel', 'eze': 'Ezekiel', 'ezk': 'Ezekiel', 'ezekiel': 'Ezekiel',
    'dan': 'Daniel', 'da': 'Daniel', 'dn': 'Daniel', 'daniel': 'Daniel',
    'hos': 'Hosea', 'ho': 'Hosea', 'hosea': 'Hosea',
    'joel': 'Joel', 'jl': 'Joel',
    'amos': 'Amos', 'am': 'Amos',
    'obad': 'Obadiah', 'ob': 'Obadiah', 'obadiah': 'Obadiah',
    'jonah': 'Jonah', 'jon': 'Jonah',
    'mic': 'Micah', 'mi': 'Micah', 'micah': 'Micah',
    'nah': 'Nahum', 'na': 'Nahum', 'nahum': 'Nahum',
    'hab': 'Habakkuk', 'hb': 'Habakkuk', 'habakkuk': 'Habakkuk',
    'zeph': 'Zephaniah', 'zep': 'Zephaniah', 'zp': 'Zephaniah', 'zephaniah': 'Zephaniah',
    'hag': 'Haggai', 'hg': 'Haggai', 'haggai': 'Haggai',
    'zech': 'Zechariah', 'zec': 'Zechariah', 'zc': 'Zechariah', 'zechariah': 'Zechariah',
    'mal': 'Malachi', 'ml': 'Malachi', 'malachi': 'Malachi',
    
    # New Testament
    'matt': 'Matthew', 'mt': 'Matthew', 'mat': 'Matthew', 'matthew': 'Matthew',
    'mark': 'Mark', 'mk': 'Mark', 'mr': 'Mark',
    'luke': 'Luke', 'lk': 'Luke', 'lu': 'Luke',
    'john': 'John', 'jn': 'John', 'jhn': 'John',
    'acts': 'Acts', 'ac': 'Acts', 'act': 'Acts',
    'rom': 'Romans', 'ro': 'Romans', 'rm': 'Romans', 'romans': 'Romans', 'rom.': 'Romans',
    '1 cor': '1 Corinthians', '1cor': '1 Corinthians', '1 co': '1 Corinthians', '1co': '1 Corinthians', '1 corinthians': '1 Corinthians',
    '2 cor': '2 Corinthians', '2cor': '2 Corinthians', '2 co': '2 Corinthians', '2co': '2 Corinthians', '2 corinthians': '2 Corinthians',
    'gal': 'Galatians', 'ga': 'Galatians', 'galatians': 'Galatians',
    'eph': 'Ephesians', 'ep': 'Ephesians', 'ephesians': 'Ephesians',
    'phil': 'Philippians', 'php': 'Philippians', 'pp': 'Philippians', 'philippians': 'Philippians',
    'col': 'Colossians', 'co': 'Colossians', 'colossians': 'Colossians',
    '1 thess': '1 Thessalonians', '1thess': '1 Thessalonians', '1 th': '1 Thessalonians', '1th': '1 Thessalonians', '1 thessalonians': '1 Thessalonians',
    '2 thess': '2 Thessalonians', '2thess': '2 Thessalonians', '2 th': '2 Thessalonians', '2th': '2 Thessalonians', '2 thessalonians': '2 Thessalonians',
    '1 tim': '1 Timothy', '1tim': '1 Timothy', '1 ti': '1 Timothy', '1ti': '1 Timothy', '1 timothy': '1 Timothy',
    '2 tim': '2 Timothy', '2tim': '2 Timothy', '2 ti': '2 Timothy', '2ti': '2 Timothy', '2 timothy': '2 Timothy',
    'titus': 'Titus', 'tit': 'Titus', 'ti': 'Titus',
    'phlm': 'Philemon', 'phm': 'Philemon', 'pm': 'Philemon', 'philemon': 'Philemon',
    'heb': 'Hebrews', 'he': 'Hebrews', 'hebrews': 'Hebrews',
    'jas': 'James', 'jm': 'James', 'ja': 'James', 'james': 'James',
    '1 pet': '1 Peter', '1pet': '1 Peter', '1 pe': '1 Peter', '1pe': '1 Peter', '1 peter': '1 Peter',
    '2 pet': '2 Peter', '2pet': '2 Peter', '2 pe': '2 Peter', '2pe': '2 Peter', '2 peter': '2 Peter',
    '1 john': '1 John', '1john': '1 John', '1 jn': '1 John', '1jn': '1 John',
    '2 john': '2 John', '2john': '2 John', '2 jn': '2 John', '2jn': '2 John',
    '3 john': '3 John', '3john': '3 John', '3 jn': '3 John', '3jn': '3 John',
    'jude': 'Jude', 'jud': 'Jude', 'jd': 'Jude',
    'rev': 'Revelation', 're': 'Revelation', 'rv': 'Revelation', 'revelation': 'Revelation'
}

def create_backup(file_path: str) -> Tuple[bool, str]:
    """
    Create a backup of the original file.
    
    Args:
        file_path: Path to the file to back up
        
    Returns:
        Tuple[bool, str]: Success status and backup file path or error message
    """
    try:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename, extension = os.path.splitext(file_path)
        backup_path = f"{filename}_backup_{timestamp}{extension}"
        
        shutil.copy2(file_path, backup_path)
        return True, backup_path
    except Exception as e:
        return False, str(e)

def compile_regex_patterns() -> List[Tuple[Pattern, str]]:
    """
    Compile regex patterns for Bible reference formats.
    Only recognizes and standardizes to "Book Chapter:Verse" format.
    
    Returns:
        List of tuples containing regex pattern and replacement format
    """
    # Create a pattern for book names
    book_pattern = r'(' + '|'.join(BIBLE_BOOKS.keys()) + r')'
    
    patterns = [
        # Format: "John 3:16" or "John 3:16-17" or "John 3:16,18"
        (re.compile(rf'\b{book_pattern}\s*(\d+):(\d+)(?:-(\d+))?(?:,\s*(\d+)(?:-(\d+))?)*\b', re.IGNORECASE), 
         lambda m: standardize_reference(m)),
        
        # Also recognize period format: "John 3.16" and convert to colon format
        (re.compile(rf'\b{book_pattern}\s*(\d+)\.(\d+)(?:-(\d+))?(?:,\s*(\d+)(?:-(\d+))?)*\b', re.IGNORECASE), 
         lambda m: standardize_reference(m)),
        
        # Also recognize "John chapter 3 verse 16" format
        (re.compile(rf'\b{book_pattern}\s*chapter\s*(\d+)\s*verse\s*(\d+)(?:-(\d+))?\b', re.IGNORECASE),
         lambda m: standardize_reference(m)),
        
        # Also recognize space-only format: "John 3 16"
        (re.compile(rf'\b{book_pattern}\s*(\d+)\s+(\d+)(?:-(\d+))?\b', re.IGNORECASE),
         lambda m: standardize_reference(m)),
         
        # Standalone chapter reference (like "Psalm 23" or "John 3")
        (re.compile(rf'\b{book_pattern}\s+(\d+)(?!\s*[:\.]\d+|\s+\d+)\b', re.IGNORECASE),
         lambda m: standardize_standalone_chapter(m)),
    ]
    
    return patterns
def standardize_standalone_chapter(match) -> str:
    """
    Standardize a standalone chapter reference (like "Psalm 23") to "Book Chapter:1" format.
    
    Args:
        match: Regex match object
        
    Returns:
        Standardized reference string in "Book Chapter:1" format
    """
    book_abbr = match.group(1).lower()
    book_name = BIBLE_BOOKS.get(book_abbr, book_abbr)  # Use the standardized name
    chapter = match.group(2)
    
    # For standalone chapter references, always append ":1"
    return f"{book_name} {chapter}:1"

def standardize_reference(match, separator: str = ':') -> str:
    """
    Standardize a Bible reference to "Book Chapter:Verse" format.
    Ensures consistent book names (expands all abbreviations) and
    no spaces after commas in verse lists.
    
    Args:
        match: Regex match object
        
    Returns:
        Standardized reference string in "Book Chapter:Verse" format
    """
    book_abbr = match.group(1).lower()
    book_name = BIBLE_BOOKS.get(book_abbr, book_abbr)  # Use the standardized name
    
    chapter = match.group(2)
    verse_start = match.group(3)
    verse_end = match.group(4) if len(match.groups()) >= 4 and match.group(4) else None
    
    # Handle additional verse ranges if present
    additional_verse_start = match.group(5) if len(match.groups()) >= 5 and match.group(5) else None
    additional_verse_end = match.group(6) if len(match.groups()) >= 6 and match.group(6) else None
    
    # Build the standardized reference - always use colon format
    reference = f"{book_name} {chapter}:{verse_start}"
    
    # Add verse range if present
    if verse_end:
        reference += f"-{verse_end}"
    
    # Add additional verses if present
    if additional_verse_start:
        # No space after comma
        reference += f",{additional_verse_start}"
        if additional_verse_end:
            reference += f"-{additional_verse_end}"
    
    return reference

def process_paragraph(paragraph: Paragraph, patterns: List[Tuple[Pattern, str]]) -> bool:
    """
    Process a single paragraph to standardize Bible references.
    
    Args:
        paragraph: Paragraph object from docx
        patterns: List of regex patterns and their replacement functions
        
    Returns:
        Boolean indicating if any changes were made
    """
    changed = False
    
    # Split the paragraph into runs to preserve formatting
    runs = list(paragraph.runs)
    if not runs:
        return False
    
    # Process each run
    for i, run in enumerate(runs):
        if run.text:
            new_text, run_changed = process_run(run.text, patterns)
            if run_changed:
                run.text = new_text
                changed = True
    
    return changed

def process_run(text: str, patterns: List[Tuple[Pattern, str]]) -> Tuple[str, bool]:
    """
    Process a single run's text to standardize Bible references.
    
    Args:
        text: Text content of the run
        patterns: List of regex patterns and their replacement functions
        
    Returns:
        Tuple of (new text, whether changes were made)
    """
    original_text = text
    changed = False
    
    # Apply each pattern to the text
    for pattern, replacement_func in patterns:
        # Use regex to find and replace all instances
        text = pattern.sub(replacement_func, text)
    
    # Additional cleanup patterns
    
    # Remove spaces after commas in verse references
    # Look for patterns like "Book 3:16, 17" and replace with "Book 3:16,17"
    text = re.sub(r'(\d+),\s+(\d+)', r'\1,\2', text)
    
    # Handle multi-chapter references with dash (e.g., "Psalm 1-2")
    # Convert to "Psalm 1:1-2:1"
    text = re.sub(r'((?:' + '|'.join(BIBLE_BOOKS.values()) + r')\s+)(\d+)-(\d+)(?!\s*:)', 
                 r'\1\2:1-\3:1', text, flags=re.IGNORECASE)
    
    # Check if changes were made
    if text != original_text:
        changed = True
    
    return text, changed

def process_document(doc_path: str, output_path: Optional[str] = None, create_backup_file: bool = True) -> Dict:
    """
    Process a Word document to standardize Bible references.
    
    Args:
        doc_path: Path to the Word document
        output_path: Path where the processed document should be saved (if None, overwrite original)
        create_backup_file: Whether to create a backup of the original file
        
    Returns:
        Dictionary with process statistics and status information
    """
    result = {
        'success': False,
        'changes_made': False,
        'paragraphs_processed': 0,
        'paragraphs_changed': 0,
        'backup_path': None,
        'output_path': None,
        'error': None
    }
    
    # Check if file exists
    if not os.path.exists(doc_path):
        result['error'] = f"File not found: {doc_path}"
        return result
    
    # Create backup if requested
    if create_backup_file:
        backup_success, backup_result = create_backup(doc_path)
        if backup_success:
            result['backup_path'] = backup_result
        else:
            result['error'] = f"Failed to create backup: {backup_result}"
            return result
    
    try:
        # Compile regex patterns
        patterns = compile_regex_patterns()
        
        # Open the document
        document = Document(doc_path)
        
        # Process each paragraph
        for paragraph in document.paragraphs:
            result['paragraphs_processed'] += 1
            
            # Process the paragraph
            if process_paragraph(paragraph, patterns):
                result['paragraphs_changed'] += 1
                result['changes_made'] = True
        
        # Process text in tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        result['paragraphs_processed'] += 1
                        if process_paragraph(paragraph, patterns):
                            result['paragraphs_changed'] += 1
                            result['changes_made'] = True
        
        # Save the document
        if output_path is None:
            output_path = doc_path
        
        document.save(output_path)
        result['output_path'] = output_path
        result['success'] = True
        
    except Exception as e:
        result['error'] = str(e)
        return result
    
    return result

def main():
    """Main entry point for the script."""
    parser = argparse.ArgumentParser(
        description="Standardize Bible references to 'Book Chapter:Verse' format in Word documents",
        epilog="""Example: python standardize_scripture.py document.docx -o standardized_document.docx
        
Note: Standalone chapter references (like 'Psalm 23') will be converted to 'Psalm 23:1'"""
    )
    
    parser.add_argument(
        "document", 
        help="Path to the Word document to process"
    )
    
    parser.add_argument(
        "-o", "--output", 
        help="Path where the processed document should be saved (if not specified, overwrites original)"
    )
    
    parser.add_argument(
        "--no-backup", 
        action="store_true", 
        help="Skip creating a backup of the original document"
    )
    
    parser.add_argument(
        "-v", "--verbose", 
        action="store_true", 
        help="Display detailed processing information"
    )
    
    args = parser.parse_args()
    
    # Process the document
    print(f"Processing {args.document}...")
    
    result = process_document(
        args.document, 
        args.output, 
        not args.no_backup
    )
    
    if result['success']:
        print(f"✓ Successfully processed document")
        
        if result['backup_path']:
            print(f"✓ Backup created at: {result['backup_path']}")
        
        if result['changes_made']:
            print(f"✓ Made changes to {result['paragraphs_changed']} out of {result['paragraphs_processed']} paragraphs")
            print(f"✓ Saved to: {result['output_path']}")
        else:
            print("ℹ No Bible references found to standardize")
        
        if args.verbose:
            print("\nDetailed statistics:")
            print(f"  - Paragraphs processed: {result['paragraphs_processed']}")
            print(f"  - Paragraphs changed: {result['paragraphs_changed']}")
            print(f"  - Change percentage: {result['paragraphs_changed']/max(result['paragraphs_processed'], 1)*100:.1f}%")
    else:
        print(f"✗ Error: {result['error']}")
        sys.exit(1)

if __name__ == "__main__":
    main()
