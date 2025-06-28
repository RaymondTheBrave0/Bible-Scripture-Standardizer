#!/usr/bin/env python3
"""
pdf_text_cleaner.py - Module for fixing common PDF text extraction spacing issues.

This module provides comprehensive text cleanup functionality specifically designed
to handle spacing problems that commonly occur when extracting text from PDF files.

Features:
- Fixes missing spaces between concatenated words
- Handles religious text-specific spacing issues
- Corrects punctuation spacing
- Cleans up Bible reference formatting
- Applies intelligent word boundary detection
"""

import re
import logging

logger = logging.getLogger(__name__)

class PDFTextCleaner:
    """
    A comprehensive text cleaner for PDF-extracted text that fixes common spacing issues.
    """
    
    def __init__(self):
        """Initialize the PDF text cleaner with predefined patterns."""
        self._compile_patterns()
    
    def _compile_patterns(self):
        """Compile regex patterns for efficient text cleaning."""
        # Common words that often get concatenated
        self.common_words = [
            'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by',
            'is', 'are', 'was', 'were', 'be', 'been', 'have', 'has', 'had', 'will', 'would',
            'could', 'should', 'may', 'might', 'can', 'do', 'does', 'did', 'this', 'that',
            'these', 'those', 'a', 'an', 'as', 'if', 'so', 'no', 'not', 'all', 'any',
            'each', 'every', 'some', 'many', 'more', 'most', 'less', 'few', 'new', 'old',
            'other', 'such', 'than', 'when', 'where', 'why', 'how', 'what', 'who', 'which',
            'whom', 'whose', 'her', 'his', 'its', 'our', 'your', 'their', 'them', 'they',
            'we', 'you', 'he', 'she', 'it', 'me', 'us', 'him'
        ]
        
        # Religious and biblical terms that commonly get concatenated
        self.religious_words = [
            'God', 'Lord', 'Jesus', 'Christ', 'Christian', 'Scripture', 'Biblical',
            'spiritual', 'holy', 'divine', 'sacred', 'blessed', 'salvation', 'redemption',
            'faith', 'believe', 'prayer', 'worship', 'church', 'temple', 'prophet',
            'apostle', 'disciple', 'gospel', 'heaven', 'eternal', 'righteous', 'sin',
            'forgive', 'grace', 'mercy', 'love', 'peace', 'joy', 'hope', 'truth',
            'light', 'life', 'death', 'resurrection', 'sacrifice', 'covenant', 'kingdom',
            'glory', 'praise', 'honor', 'almighty', 'creator', 'savior', 'redeemer',
            'messiah', 'trinity', 'father', 'son', 'spirit', 'angel', 'miracle',
            'parable', 'psalm', 'proverb', 'revelation', 'prophecy', 'testament',
            'epistle', 'commandment', 'blessing', 'sanctify', 'consecrate', 'minister',
            'pastor', 'priest', 'deacon', 'elder', 'bishop'
        ]
        
        # Prepositions and articles that commonly get concatenated
        self.prepositions = [
            'the', 'a', 'an', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by',
            'from', 'into', 'unto', 'through', 'over', 'under', 'above', 'below',
            'before', 'after', 'during', 'within', 'without', 'between', 'among'
        ]
    
    def fix_text_spacing(self, text: str) -> str:
        """
        Fix common spacing issues in extracted PDF text.
        
        Args:
            text: Input text with potential spacing issues
            
        Returns:
            Text with improved spacing
        """
        if not text or not text.strip():
            return text
        
        logger.debug(f"Cleaning text: {text[:50]}...")
        original_text = text
        
        # Apply cleaning steps in order
        text = self._fix_basic_concatenations(text)
        text = self._fix_punctuation_spacing(text)
        text = self._fix_common_word_concatenations(text)
        text = self._fix_religious_word_concatenations(text)
        text = self._fix_bible_reference_spacing(text)
        text = self._fix_sentence_boundaries(text)
        text = self._fix_preposition_concatenations(text)
        text = self._cleanup_multiple_spaces(text)
        
        if text != original_text:
            logger.debug(f"Text cleaning applied: '{original_text[:30]}...' -> '{text[:30]}...'")
        
        return text
    
    def _fix_basic_concatenations(self, text: str) -> str:
        """Fix basic word concatenation patterns - very conservative approach."""
        # Only fix very obvious patterns to avoid breaking legitimate text
        
        # Fix digit+letter and letter+digit (less likely to break real words)
        text = re.sub(r'([0-9])([a-zA-Z])', r'\1 \2', text)
        text = re.sub(r'([a-z])([0-9])', r'\1 \2', text)
        
        # Only fix camelCase if it looks like two complete words
        # This is very conservative - we'll let other patterns handle most cases
        
        return text
    
    def _fix_punctuation_spacing(self, text: str) -> str:
        """Fix spacing around punctuation marks."""
        # Fix missing spaces after punctuation
        text = re.sub(r'([.!?])([A-Z])', r'\1 \2', text)
        text = re.sub(r'([,;:])([a-zA-Z])', r'\1 \2', text)
        
        # Fix missing spaces before quotation marks
        text = re.sub(r'([a-zA-Z])(["\'])', r'\1 \2', text)
        
        return text
    
    def _fix_common_word_concatenations(self, text: str) -> str:
        """Fix concatenations involving common English words."""
        # Only fix obvious concatenations - be more conservative
        # Look for specific patterns where a common word is clearly concatenated
        
        # Fix patterns like "butchosen", "andprecious", etc.
        # Only if the word is 3+ characters to avoid breaking legitimate compound words
        specific_fixes = [
            (r'\bbut([a-z]{3,})', r'but \1'),
            (r'\band([a-z]{3,})', r'and \1'),
            (r'\bor([a-z]{3,})', r'or \1'),
            (r'\bin([a-z]{3,})', r'in \1'),
            (r'\bof([a-z]{3,})', r'of \1'),
            (r'\bto([a-z]{3,})', r'to \1'),
            (r'\bfor([a-z]{3,})', r'for \1'),
            (r'\bwith([a-z]{3,})', r'with \1'),
            (r'\bthe([a-z]{3,})', r'the \1'),
            (r'\bthat([a-z]{3,})', r'that \1'),
            (r'\bthis([a-z]{3,})', r'this \1'),
            (r'\bwho([a-z]{3,})', r'who \1'),
            (r'\bwhen([a-z]{3,})', r'when \1'),
            (r'\bwhere([a-z]{3,})', r'where \1'),
        ]
        
        for pattern, replacement in specific_fixes:
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        
        return text
    
    def _fix_religious_word_concatenations(self, text: str) -> str:
        """Fix concatenations specific to religious texts."""
        # Be more conservative - only fix obvious religious concatenations
        specific_religious_fixes = [
            (r'\bpriceless([a-z]{2,})', r'priceless \1'),
            (r'\bspiritual([a-z]{3,})', r'spiritual \1'),
            (r'\bChristian([a-z]{3,})', r'Christian \1'),
            (r'\bScripture([a-z]{3,})', r'Scripture \1'),
            (r'\bBiblical([a-z]{3,})', r'Biblical \1'),
            (r'\bsacred([a-z]{3,})', r'sacred \1'),
            (r'\bblessed([a-z]{3,})', r'blessed \1'),
            (r'\beternal([a-z]{3,})', r'eternal \1'),
            (r'\brighteous([a-z]{3,})', r'righteous \1'),
            (r'\bsalvation([a-z]{3,})', r'salvation \1'),
            (r'\bredemption([a-z]{3,})', r'redemption \1'),
            (r'\bresurrection([a-z]{3,})', r'resurrection \1'),
        ]
        
        for pattern, replacement in specific_religious_fixes:
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        
        return text
    
    def _fix_bible_reference_spacing(self, text: str) -> str:
        """Fix spacing issues specific to Bible references."""
        # Fix book names concatenated with references
        text = re.sub(r'([a-zA-Z])([0-9]+:[0-9]+)', r'\1 \2', text)
        text = re.sub(r'([0-9]+:[0-9]+)([a-zA-Z])', r'\1 \2', text)
        
        # Fix common Bible book abbreviations
        bible_books = ['Gen', 'Ex', 'Lev', 'Num', 'Deut', 'Josh', 'Judg', 'Ruth', 'Sam', 'Kings', 'Chr', 'Ezra', 'Neh', 'Esth', 'Job', 'Ps', 'Prov', 'Eccl', 'Song', 'Isa', 'Jer', 'Lam', 'Ezek', 'Dan', 'Hos', 'Joel', 'Amos', 'Obad', 'Jonah', 'Mic', 'Nah', 'Hab', 'Zeph', 'Hag', 'Zech', 'Mal', 'Mt', 'Mk', 'Lk', 'Jn', 'Acts', 'Rom', 'Cor', 'Gal', 'Eph', 'Phil', 'Col', 'Thess', 'Tim', 'Tit', 'Phlm', 'Heb', 'Jas', 'Pet', 'Rev']
        
        for book in bible_books:
            text = re.sub(rf'([a-z])({book})\b', r'\1 \2', text, flags=re.IGNORECASE)
            text = re.sub(rf'\b({book})([a-z])', r'\1 \2', text, flags=re.IGNORECASE)
        
        return text
    
    def _fix_sentence_boundaries(self, text: str) -> str:
        """Fix concatenated words at sentence boundaries."""
        # Fix words concatenated at sentence starts
        text = re.sub(r'([a-z])([A-Z][a-z])', r'\1 \2', text)
        
        # Fix period followed immediately by capital letter without space
        text = re.sub(r'\.([A-Z])', r'. \1', text)
        
        return text
    
    def _fix_preposition_concatenations(self, text: str) -> str:
        """Fix concatenations involving prepositions and articles."""
        prep_pattern = '|'.join(self.prepositions)
        
        # Fix patterns like: word+preposition+word
        text = re.sub(
            rf'([a-z])({prep_pattern})([A-Z])',
            r'\1 \2 \3',
            text,
            flags=re.IGNORECASE
        )
        
        return text
    
    def _cleanup_multiple_spaces(self, text: str) -> str:
        """Clean up multiple consecutive spaces."""
        # Replace multiple spaces with single space
        text = re.sub(r'\s+', ' ', text)
        
        # Clean up spaces around punctuation
        text = re.sub(r'\s+([.!?,:;])', r'\1', text)
        text = re.sub(r'([.!?])\s+', r'\1 ', text)
        
        return text.strip()
    
    def clean_document_paragraphs(self, doc_path: str) -> tuple[bool, int]:
        """
        Apply text cleaning to all paragraphs in a DOCX document.
        
        Args:
            doc_path: Path to the DOCX document
            
        Returns:
            Tuple of (success, paragraphs_cleaned)
        """
        try:
            from docx import Document
            
            doc = Document(doc_path)
            paragraphs_cleaned = 0
            
            # Clean main document paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    original_text = paragraph.text
                    cleaned_text = self.fix_text_spacing(original_text)
                    
                    if cleaned_text != original_text:
                        paragraph.clear()
                        paragraph.add_run(cleaned_text)
                        paragraphs_cleaned += 1
            
            # Clean text in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                original_text = paragraph.text
                                cleaned_text = self.fix_text_spacing(original_text)
                                
                                if cleaned_text != original_text:
                                    paragraph.clear()
                                    paragraph.add_run(cleaned_text)
                                    paragraphs_cleaned += 1
            
            # Save the cleaned document
            doc.save(doc_path)
            
            if paragraphs_cleaned > 0:
                logger.info(f"Applied text spacing cleanup to {paragraphs_cleaned} paragraphs")
            
            return True, paragraphs_cleaned
            
        except Exception as e:
            logger.error(f"Failed to clean document paragraphs: {str(e)}")
            return False, 0


# Convenience functions for easy use
def clean_pdf_text(text: str) -> str:
    """
    Convenience function to clean PDF-extracted text.
    
    Args:
        text: Text with potential spacing issues
        
    Returns:
        Cleaned text
    """
    cleaner = PDFTextCleaner()
    return cleaner.fix_text_spacing(text)


def clean_pdf_document(doc_path: str) -> tuple[bool, int]:
    """
    Convenience function to clean all text in a DOCX document.
    
    Args:
        doc_path: Path to the DOCX document
        
    Returns:
        Tuple of (success, paragraphs_cleaned)
    """
    cleaner = PDFTextCleaner()
    return cleaner.clean_document_paragraphs(doc_path)
