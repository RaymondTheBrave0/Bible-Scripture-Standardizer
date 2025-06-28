#!/usr/bin/env python3
"""
Create multiple test documents for directory processing demo.
"""

from docx import Document
import os

def create_test_documents():
    """Create multiple test documents with different Bible reference formats."""
    
    # Ensure test_documents directory exists
    os.makedirs('test_documents', exist_ok=True)
    
    # Document 1: Sermon Notes
    doc1 = Document()
    doc1.add_heading('Sunday Sermon Notes', 0)
    doc1.add_paragraph("Today we studied gen 1:1 and how God created everything.")
    doc1.add_paragraph("The pastor also referenced jn 3.16 about God's love.")
    doc1.add_paragraph("We discussed ps 23 and David's trust in God.")
    doc1.save('test_documents/sermon_notes.docx')
    
    # Document 2: Bible Study
    doc2 = Document()
    doc2.add_heading('Bible Study Group - Week 3', 0)
    doc2.add_paragraph("This week we're covering matt 5:3-12, the beatitudes.")
    doc2.add_paragraph("Cross-reference with luke 6 9-23 for Luke's version.")
    doc2.add_paragraph("Also read 1 cor 13:4-8 about love.")
    doc2.save('test_documents/bible_study.docx')
    
    # Document 3: Personal Devotions
    doc3 = Document()
    doc3.add_heading('Personal Devotions', 0)
    doc3.add_paragraph("Morning reading: isa chapter 53 verse 5.")
    doc3.add_paragraph("Evening prayer based on phil 4.13.")
    doc3.add_paragraph("Memory verse: rom 8.28 - all things work together.")
    doc3.save('test_documents/devotions.docx')
    
    print("Created 3 test documents in test_documents/ directory:")
    print("  - sermon_notes.docx")
    print("  - bible_study.docx") 
    print("  - devotions.docx")

if __name__ == "__main__":
    create_test_documents()
