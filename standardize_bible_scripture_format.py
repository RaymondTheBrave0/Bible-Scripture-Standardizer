#!/usr/bin/env python3
"""
standardize_bible_scripture_format.py - Production-ready module for standardizing Bible references in Word documents.

This module provides comprehensive Bible reference standardization using a CSV file containing
all possible abbreviations and formats. It handles various edge cases and provides robust
error handling suitable for production environments.

Features:
- Loads Bible book mappings from CSV file
- Handles multiple abbreviation formats per book
- Supports various reference formats (colon, period, space, etc.)
- Preserves document formatting while standardizing references
- Creates backups before modification
- Comprehensive logging and error handling
- Production-ready with proper testing support

Dependencies:
    - python-docx: pip install python-docx
    - Standard library modules: re, os, csv, logging, datetime, shutil
"""

import os
import re
import csv
import logging
import shutil
from datetime import datetime
from typing import Dict, List, Tuple, Optional, Pattern
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class BibleReferenceStandardizer:
    """
    A comprehensive Bible reference standardizer that loads book mappings from CSV
    and provides robust standardization capabilities.
    """
    
    def __init__(self, csv_path: str = None):
        """
        Initialize the standardizer with Bible book mappings.
        
        Args:
            csv_path: Path to the CSV file containing Bible book abbreviations
        """
        if csv_path is None:
            csv_path = os.path.join(os.path.dirname(__file__), 'Bible-Books-Abbr.csv')
        
        self.bible_books = {}
        self.book_variations = {}
        self.patterns = []
        
        self._load_bible_books(csv_path)
        self._compile_patterns()
        
        logger.info(f"Loaded {len(self.bible_books)} Bible books with {len(self.book_variations)} total variations")
    
    def _load_bible_books(self, csv_path: str) -> None:
        """
        Load Bible book names and abbreviations from CSV file.
        
        Args:
            csv_path: Path to the CSV file
        """
        if not os.path.exists(csv_path):
            raise FileNotFoundError(f"Bible books CSV file not found: {csv_path}")
        
        try:
            with open(csv_path, 'r', encoding='utf-8') as file:
                reader = csv.reader(file)
                for row_num, row in enumerate(reader, 1):
                    if len(row) < 2:
                        continue
                    
                    full_name = row[0].strip()
                    abbreviations = row[1].strip() if len(row) > 1 else ''
                    
                    if not full_name:
                        continue
                    
                    # Store the canonical name (clean up parenthetical info)
                    clean_name = full_name
                    if '(' in clean_name:
                        clean_name = clean_name.split('(')[0].strip()
                    self.bible_books[full_name.lower()] = clean_name
                    
                    # Parse abbreviations (they may be comma-separated in quotes)
                    if abbreviations:
                        # Remove quotes and split by comma
                        abbr_list = [abbr.strip() for abbr in abbreviations.strip('"').split(',')]
                        
                        for abbr in abbr_list:
                            abbr = abbr.strip()
                            if abbr:
                                # Handle special cases like "Ps (pl. Pss)"
                                if '(' in abbr:
                                    # Extract both the base and the parenthetical form
                                    base_abbr = abbr.split('(')[0].strip()
                                    if base_abbr:
                                        self.book_variations[base_abbr.lower()] = clean_name
                                    
                                    # Extract content within parentheses
                                    paren_content = abbr[abbr.find('(')+1:abbr.find(')')]
                                    if 'pl.' in paren_content:
                                        plural_form = paren_content.replace('pl.', '').strip()
                                        if plural_form:
                                            self.book_variations[plural_form.lower()] = clean_name
                                else:
                                    self.book_variations[abbr.lower()] = clean_name
                    
                    # Also add the full name as a variation (both original and clean)
                    self.book_variations[full_name.lower()] = clean_name
                    self.book_variations[clean_name.lower()] = clean_name
                    
                    # Special handling for Psalm/Psalms
                    if clean_name == 'Psalms':
                        self.book_variations['psalm'] = clean_name
                        
        except Exception as e:
            raise Exception(f"Error loading Bible books from CSV: {str(e)}")
    
    def _compile_patterns(self) -> None:
        """
        Compile regex patterns for various Bible reference formats.
        """
        # Create a pattern for all book variations
        book_variations_escaped = [re.escape(book) for book in self.book_variations.keys()]
        # Sort by length (longest first) to match longer abbreviations first
        book_variations_escaped.sort(key=len, reverse=True)
        book_pattern = r'(' + '|'.join(book_variations_escaped) + r')'
        
        self.patterns = [
            # Standard format: "John 3:16" or "John 3:16-17" or "John 3:16,18"
            (re.compile(rf'\b{book_pattern}\s+(\d+):(\d+)(?:-(\d+))?(?:,\s*(\d+)(?:-(\d+))?)*\b', re.IGNORECASE),
             self._standardize_reference),
            
            # Period format: "John 3.16" -> "John 3:16"
            (re.compile(rf'\b{book_pattern}\s+(\d+)\.(\d+)(?:-(\d+))?(?:,\s*(\d+)(?:-(\d+))?)*\b', re.IGNORECASE),
             self._standardize_reference),
            
            # Space format: "John 3 16" -> "John 3:16"
            (re.compile(rf'\b{book_pattern}\s+(\d+)\s+(\d+)(?:-(\d+))?\b', re.IGNORECASE),
             self._standardize_reference),
            
            # Chapter-verse format: "John chapter 3 verse 16" -> "John 3:16"
            (re.compile(rf'\b{book_pattern}\s*chapter\s+(\d+)\s*verse\s+(\d+)(?:-(\d+))?\b', re.IGNORECASE),
             self._standardize_reference),
            
            # Standalone chapter: "Psalm 23" -> "Psalm 23:1" (only for single-chapter books or when no verse specified)
            (re.compile(rf'\b{book_pattern}\s+(\d+)(?!\s*[:\.\d]|\s+\d+)\b', re.IGNORECASE),
             self._standardize_chapter_only),
            
            # Handle verse ranges across chapters: "John 3:16-4:2"
            (re.compile(rf'\b{book_pattern}\s+(\d+):(\d+)-(\d+):(\d+)\b', re.IGNORECASE),
             self._standardize_cross_chapter_range),
        ]
    
    def _standardize_reference(self, match) -> str:
        """
        Standardize a Bible reference to "Book Chapter:Verse" format.
        
        Args:
            match: Regex match object
            
        Returns:
            Standardized reference string
        """
        book_abbr = match.group(1).lower()
        book_name = self.book_variations.get(book_abbr, match.group(1))
        
        chapter = match.group(2)
        verse_start = match.group(3)
        
        # Build the base reference
        reference = f"{book_name} {chapter}:{verse_start}"
        
        # Handle verse ranges and additional verses
        groups = match.groups()
        if len(groups) >= 4 and groups[3]:  # verse end range
            reference += f"-{groups[3]}"
        
        if len(groups) >= 5 and groups[4]:  # additional verse start
            reference += f",{groups[4]}"
            if len(groups) >= 6 and groups[5]:  # additional verse end
                reference += f"-{groups[5]}"
        
        return reference
    
    def _standardize_chapter_only(self, match) -> str:
        """
        Standardize a chapter-only reference like "Psalm 23" to "Psalm 23:1".
        
        Args:
            match: Regex match object
            
        Returns:
            Standardized reference string
        """
        book_abbr = match.group(1).lower()
        book_name = self.book_variations.get(book_abbr, match.group(1))
        chapter = match.group(2)
        
        return f"{book_name} {chapter}:1"
    
    def _standardize_cross_chapter_range(self, match) -> str:
        """
        Standardize cross-chapter verse ranges like "John 3:16-4:2".
        
        Args:
            match: Regex match object
            
        Returns:
            Standardized reference string
        """
        book_abbr = match.group(1).lower()
        book_name = self.book_variations.get(book_abbr, match.group(1))
        
        start_chapter = match.group(2)
        start_verse = match.group(3)
        end_chapter = match.group(4)
        end_verse = match.group(5)
        
        return f"{book_name} {start_chapter}:{start_verse}-{end_chapter}:{end_verse}"
    
    # Note: Complex text spacing functionality moved to pdf_text_cleaner.py module
    
    def process_text(self, text: str) -> Tuple[str, bool]:
        """
        Process text to standardize Bible references.
        
        Args:
            text: Input text
            
        Returns:
            Tuple of (processed text, whether changes were made)
        """
        original_text = text
        
        # Apply each pattern
        for pattern, replacement_func in self.patterns:
            text = pattern.sub(replacement_func, text)
        
        # Additional cleanup
        # Remove extra spaces after commas in verse lists
        text = re.sub(r'(\d+),\s+(\d+)', r'\1,\2', text)
        
        # Handle multi-chapter references with dash (e.g., "Psalm 1-3")
        book_names_escaped = [re.escape(name) for name in self.bible_books.values()]
        book_names_pattern = '|'.join(book_names_escaped)
        text = re.sub(
            rf'\b({book_names_pattern})\s+(\d+)-(\d+)(?!\s*:)',
            r'\1 \2:1-\3:1',
            text,
            flags=re.IGNORECASE
        )
        
        changed = text != original_text
        return text, changed
    
    def process_paragraph(self, paragraph: Paragraph) -> bool:
        """
        Process a single paragraph to standardize Bible references.
        
        Args:
            paragraph: Paragraph object from docx
            
        Returns:
            Boolean indicating if any changes were made
        """
        changed = False
        runs = list(paragraph.runs)
        
        if not runs:
            return False
        
        for run in runs:
            if run.text:
                new_text, run_changed = self.process_text(run.text)
                if run_changed:
                    run.text = new_text
                    changed = True
        
        return changed
    
    def create_backup(self, file_path: str) -> Tuple[bool, str]:
        """
        Create a backup of the original file.
        
        Args:
            file_path: Path to the file to back up
            
        Returns:
            Tuple[bool, str]: Success status and backup file path or error message
        """
        try:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename, extension = os.path.splitext(file_path)
            backup_path = f"{filename}_backup_{timestamp}{extension}"
            
            shutil.copy2(file_path, backup_path)
            logger.info(f"Backup created: {backup_path}")
            return True, backup_path
        except Exception as e:
            logger.error(f"Failed to create backup: {str(e)}")
            return False, str(e)
    
    def process_document(self, doc_path: str, output_path: Optional[str] = None, 
                        create_backup: bool = True) -> Dict:
        """
        Process a Word document to standardize Bible references.
        
        Args:
            doc_path: Path to the Word document
            output_path: Path where processed document should be saved (None to overwrite)
            create_backup: Whether to create a backup before processing
            
        Returns:
            Dictionary with processing results and statistics
        """
        result = {
            'success': False,
            'changes_made': False,
            'paragraphs_processed': 0,
            'paragraphs_changed': 0,
            'backup_path': None,
            'output_path': doc_path if output_path is None else output_path,
            'error': None,
            'processing_time': 0
        }
        
        start_time = datetime.now()
        
        try:
            # Check if file exists
            if not os.path.exists(doc_path):
                result['error'] = f"File not found: {doc_path}"
                logger.error(result['error'])
                return result
            
            # Create backup if requested
            if create_backup:
                backup_success, backup_result = self.create_backup(doc_path)
                if backup_success:
                    result['backup_path'] = backup_result
                else:
                    result['error'] = f"Failed to create backup: {backup_result}"
                    return result
            
            logger.info(f"Processing document: {doc_path}")
            
            # Open the document
            document = Document(doc_path)
            
            # Process paragraphs
            for paragraph in document.paragraphs:
                result['paragraphs_processed'] += 1
                if self.process_paragraph(paragraph):
                    result['paragraphs_changed'] += 1
                    result['changes_made'] = True
            
            # Process tables
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            result['paragraphs_processed'] += 1
                            if self.process_paragraph(paragraph):
                                result['paragraphs_changed'] += 1
                                result['changes_made'] = True
            
            # Save the document
            document.save(result['output_path'])
            result['success'] = True
            
            # Calculate processing time
            end_time = datetime.now()
            result['processing_time'] = (end_time - start_time).total_seconds()
            
            logger.info(f"Document processed successfully. Changes made to {result['paragraphs_changed']} paragraphs.")
            
        except Exception as e:
            result['error'] = str(e)
            logger.error(f"Error processing document: {str(e)}")
        
        return result


# Multi-format file processing functions
def process_html_file(file_path: str, standardizer) -> Dict:
    """
    Process an HTML file to standardize Bible references.
    
    Args:
        file_path: Path to the HTML file
        standardizer: BibleReferenceStandardizer instance
        
    Returns:
        Dictionary with process statistics and status information
    """
    result = {
        'success': False,
        'changes_made': False,
        'elements_processed': 0,
        'elements_changed': 0,
        'backup_path': None,
        'output_path': file_path,
        'error': None,
        'processing_time': 0
    }
    
    start_time = datetime.now()
    
    try:
        # Import HTML processing libraries with fallback
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            result['error'] = "BeautifulSoup4 not installed. Install with: pip install beautifulsoup4"
            return result
        
        # Backup the original file
        backup_success, backup_path = standardizer.create_backup(file_path)
        if not backup_success:
            result['error'] = f"Failed to create backup: {backup_path}"
            return result
        result['backup_path'] = backup_path
        
        # Read and parse HTML
        with open(file_path, 'r', encoding='utf-8') as html_file:
            content = html_file.read()
            soup = BeautifulSoup(content, 'html.parser')
        
        # Process text nodes in the HTML
        changed = False
        for text_node in soup.find_all(text=True):
            if text_node.parent.name not in ['script', 'style']:  # Skip script and style tags
                original_text = str(text_node)
                new_text, was_changed = standardizer.process_text(original_text)
                result['elements_processed'] += 1
                
                if was_changed:
                    text_node.replace_with(new_text)
                    changed = True
                    result['elements_changed'] += 1
        
        # Save changes if any
        if changed:
            with open(file_path, 'w', encoding='utf-8') as out_file:
                out_file.write(str(soup))
            result['changes_made'] = True
        
        result['success'] = True
        
    except Exception as e:
        result['error'] = str(e)
        logger.error(f"Error processing HTML file: {str(e)}")
    
    # Calculate processing time
    end_time = datetime.now()
    result['processing_time'] = (end_time - start_time).total_seconds()
    
    return result


def process_pdf_file(file_path: str, standardizer) -> Dict:
    """
    Process a PDF file to standardize Bible references.
    Converts PDF to DOCX first to preserve formatting, then processes normally.
    
    Args:
        file_path: Path to the PDF file
        standardizer: BibleReferenceStandardizer instance
        
    Returns:
        Dictionary with process statistics and status information
    """
    result = {
        'success': False,
        'changes_made': False,
        'paragraphs_processed': 0,
        'paragraphs_changed': 0,
        'backup_path': None,
        'output_path': file_path,
        'error': None,
        'processing_time': 0,
        'docx_path': None,
        'conversion_method': None
    }
    
    start_time = datetime.now()
    
    try:
        # Create backup of original PDF
        backup_success, backup_path = standardizer.create_backup(file_path)
        if not backup_success:
            result['error'] = f"Failed to create backup: {backup_path}"
            return result
        result['backup_path'] = backup_path
        
        # Step 1: Convert PDF to DOCX
        docx_path = file_path.replace('.pdf', '_converted.docx')
        conversion_success = False
        
        # Try pdftotext first (excellent text spacing, professional tool)
        try:
            import subprocess
            from docx import Document
            
            # Use pdftotext to extract text with proper spacing
            cmd = ['pdftotext', '-layout', file_path, '-']
            result_subprocess = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            
            if result_subprocess.returncode == 0 and result_subprocess.stdout.strip():
                extracted_text = result_subprocess.stdout
                
                # Create DOCX from the well-spaced text
                word_doc = Document()
                
                # Split into paragraphs and add to document
                paragraphs = extracted_text.split('\n\n')
                for para_text in paragraphs:
                    cleaned_para = para_text.strip()
                    if cleaned_para:
                        # Remove excessive line breaks but preserve intentional formatting
                        cleaned_para = ' '.join(cleaned_para.split())
                        word_doc.add_paragraph(cleaned_para)
                
                word_doc.save(docx_path)
                
                conversion_success = True
                result['conversion_method'] = 'pdftotext (professional quality - excellent spacing)'
                logger.info(f"Successfully converted PDF to DOCX using pdftotext: {docx_path}")
            else:
                logger.warning(f"pdftotext extraction failed or returned empty result")
                
        except FileNotFoundError:
            logger.warning("pdftotext not found, trying alternative method...")
        except subprocess.TimeoutExpired:
            logger.warning("pdftotext conversion timed out, trying alternative method...")
        except Exception as e:
            logger.warning(f"pdftotext conversion failed: {str(e)}, trying alternative method...")
        
        # Fallback: Try LibreOffice headless (good quality - same engine as GUI)
        if not conversion_success:
            try:
                import subprocess
                import tempfile
                import os
                
                # Create temporary directory for conversion
                with tempfile.TemporaryDirectory() as temp_dir:
                    # Use LibreOffice headless to convert PDF to DOCX
                    cmd = [
                        'libreoffice',
                        '--headless',
                        '--convert-to', 'docx',
                        '--outdir', temp_dir,
                        file_path
                    ]
                    
                    result_subprocess = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                    
                    if result_subprocess.returncode == 0:
                        # Find the converted file
                        base_name = os.path.splitext(os.path.basename(file_path))[0]
                        temp_docx = os.path.join(temp_dir, f"{base_name}.docx")
                        
                        if os.path.exists(temp_docx):
                            # Move to our target location
                            import shutil
                            shutil.move(temp_docx, docx_path)
                            
                            conversion_success = True
                            result['conversion_method'] = 'LibreOffice (professional quality)'
                            logger.info(f"Successfully converted PDF to DOCX using LibreOffice: {docx_path}")
                        else:
                            logger.warning("LibreOffice conversion completed but output file not found")
                    else:
                        logger.warning(f"LibreOffice conversion failed: {result_subprocess.stderr}")
                        
            except subprocess.TimeoutExpired:
                logger.warning("LibreOffice conversion timed out, trying alternative method...")
            except FileNotFoundError:
                logger.warning("LibreOffice not found, trying alternative method...")
            except Exception as e:
                logger.warning(f"LibreOffice conversion failed: {str(e)}, trying alternative method...")
        
        # Fallback: Try pdf2docx (good quality)
        if not conversion_success:
            try:
                from pdf2docx import Converter
                
                cv = Converter(file_path)
                cv.convert(docx_path, start=0, end=None)
                cv.close()
                
                conversion_success = True
                result['conversion_method'] = 'pdf2docx (high quality)'
                logger.info(f"Successfully converted PDF to DOCX using pdf2docx: {docx_path}")
                
            except ImportError:
                logger.warning("pdf2docx not available, trying alternative method...")
            except Exception as e:
                logger.warning(f"pdf2docx conversion failed: {str(e)}, trying alternative method...")
        
        # Fallback: Try PyMuPDF + python-docx (good formatting preservation)
        if not conversion_success:
            try:
                import fitz  # PyMuPDF
                from docx import Document
                from docx.shared import Inches
                
                # Extract text with formatting info from PDF
                pdf_doc = fitz.open(file_path)
                word_doc = Document()
                
                for page_num in range(len(pdf_doc)):
                    page = pdf_doc[page_num]
                    
                    # Extract text blocks with better spacing detection
                    blocks = page.get_text("dict")
                    
                    for block in blocks["blocks"]:
                        if "lines" in block:
                            block_text = ""
                            
                            for line in block["lines"]:
                                line_text = ""
                                prev_span = None
                                
                                for span in line["spans"]:
                                    span_text = span["text"]
                                    
                                    # Add space detection between spans
                                    if prev_span and span_text:
                                        # Calculate horizontal distance between spans
                                        prev_x1 = prev_span["bbox"][2]  # right edge of previous span
                                        curr_x0 = span["bbox"][0]       # left edge of current span
                                        gap = curr_x0 - prev_x1
                                        
                                        # If there's a significant gap, add a space
                                        if gap > 1.0:  # threshold for word spacing
                                            line_text += " "
                                    
                                    line_text += span_text
                                    prev_span = span
                                
                                # Clean up the line text and add to block
                                if line_text.strip():
                                    block_text += line_text.strip() + " "
                            
                            # Additional text cleanup using PDF text cleaner
                            if block_text.strip():
                                # Import and use the PDF text cleaner
                                from pdf_text_cleaner import clean_pdf_text
                                cleaned_text = clean_pdf_text(block_text.strip())
                                p = word_doc.add_paragraph(cleaned_text)
                    
                    # Add page break except for last page
                    if page_num < len(pdf_doc) - 1:
                        word_doc.add_page_break()
                
                pdf_doc.close()
                word_doc.save(docx_path)
                
                conversion_success = True
                result['conversion_method'] = 'PyMuPDF + python-docx (good quality)'
                logger.info(f"Successfully converted PDF to DOCX using PyMuPDF: {docx_path}")
                
            except ImportError:
                logger.warning("PyMuPDF not available, trying basic conversion...")
            except Exception as e:
                logger.warning(f"PyMuPDF conversion failed: {str(e)}, trying basic conversion...")
        
        # Last resort: Basic text extraction + simple DOCX creation
        if not conversion_success:
            try:
                import PyPDF2
                from docx import Document
                
                # Extract text using PyPDF2
                extracted_text = ""
                with open(file_path, 'rb') as pdf_file:
                    reader = PyPDF2.PdfReader(pdf_file)
                    for page in reader.pages:
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                extracted_text += page_text + "\n\n"
                        except Exception as e:
                            logger.warning(f"Could not extract text from page: {str(e)}")
                
                # Create basic DOCX
                word_doc = Document()
                # Split into paragraphs and add to document
                paragraphs = extracted_text.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        word_doc.add_paragraph(para_text.strip())
                
                word_doc.save(docx_path)
                
                conversion_success = True
                result['conversion_method'] = 'PyPDF2 + python-docx (basic quality)'
                logger.info(f"Successfully converted PDF to DOCX using PyPDF2: {docx_path}")
                
            except ImportError:
                result['error'] = "No PDF processing libraries available. Install with: pip install pdf2docx PyMuPDF or PyPDF2"
                return result
            except Exception as e:
                result['error'] = f"All PDF conversion methods failed: {str(e)}"
                return result
        
        if not conversion_success:
            result['error'] = "Failed to convert PDF to DOCX format"
            return result
        
        result['docx_path'] = docx_path
        
        # Step 2: Apply general text spacing cleanup to the converted DOCX (skip for pdftotext)
        # pdftotext already produces excellent spacing, so we skip cleanup for it
        if result['conversion_method'] != 'pdftotext (professional quality - excellent spacing)':
            try:
                from pdf_text_cleaner import clean_pdf_document
                
                cleanup_success, paragraphs_cleaned = clean_pdf_document(docx_path)
                if cleanup_success and paragraphs_cleaned > 0:
                    logger.info(f"Applied text spacing cleanup to {paragraphs_cleaned} paragraphs in converted DOCX")
                else:
                    logger.info("No text spacing issues detected in converted DOCX")
                    
            except Exception as e:
                logger.warning(f"Text spacing cleanup failed: {str(e)}, proceeding with original conversion...")
        else:
            logger.info("Skipping text spacing cleanup for pdftotext conversion (excellent spacing already)")
        
        # Step 3: Process the converted DOCX file normally for Bible reference standardization
        logger.info(f"Processing converted DOCX file: {docx_path}")
        
        # Use the existing DOCX processing logic
        docx_result = standardizer.process_document(docx_path, None, False)  # No backup needed, we have the original
        
        # Merge results
        result['success'] = docx_result['success']
        result['changes_made'] = docx_result['changes_made']
        result['paragraphs_processed'] = docx_result['paragraphs_processed']
        result['paragraphs_changed'] = docx_result['paragraphs_changed']
        
        if docx_result['success']:
            # Replace the original PDF's extension for final output (use original name)
            final_docx_path = file_path.replace('.pdf', '.docx')
            
            # Move the processed DOCX to final location
            import shutil
            shutil.move(docx_path, final_docx_path)
            result['output_path'] = final_docx_path
            
            logger.info(f"PDF converted and standardized. DOCX output saved to: {final_docx_path}")
        else:
            result['error'] = docx_result.get('error', 'Unknown error during DOCX processing')
            # Clean up temporary DOCX file on error
            try:
                os.remove(docx_path)
            except:
                pass
        
    except Exception as e:
        result['error'] = str(e)
        logger.error(f"Error processing PDF file: {str(e)}")
        # Clean up any temporary files
        try:
            docx_path = file_path.replace('.pdf', '_converted.docx')
            if os.path.exists(docx_path):
                os.remove(docx_path)
        except:
            pass
    
    # Calculate processing time
    end_time = datetime.now()
    result['processing_time'] = (end_time - start_time).total_seconds()
    
    return result


def process_txt_file(file_path: str, standardizer) -> Dict:
    """
    Process a plain text file to standardize Bible references.
    
    Args:
        file_path: Path to the text file
        standardizer: BibleReferenceStandardizer instance
        
    Returns:
        Dictionary with process statistics and status information
    """
    result = {
        'success': False,
        'changes_made': False,
        'lines_processed': 0,
        'lines_changed': 0,
        'backup_path': None,
        'output_path': file_path,
        'error': None,
        'processing_time': 0
    }
    
    start_time = datetime.now()
    
    try:
        # Backup the original file
        backup_success, backup_path = standardizer.create_backup(file_path)
        if not backup_success:
            result['error'] = f"Failed to create backup: {backup_path}"
            return result
        result['backup_path'] = backup_path
        
        # Read file
        with open(file_path, 'r', encoding='utf-8') as txt_file:
            lines = txt_file.readlines()
        
        # Process each line
        changed = False
        for i, line in enumerate(lines):
            new_line, line_changed = standardizer.process_text(line)
            result['lines_processed'] += 1
            if line_changed:
                lines[i] = new_line
                changed = True
                result['lines_changed'] += 1
        
        # Save the file if any changes were made
        if changed:
            with open(file_path, 'w', encoding='utf-8') as txt_file:
                txt_file.writelines(lines)
            result['changes_made'] = True
        
        result['success'] = True
        
    except Exception as e:
        result['error'] = str(e)
        logger.error(f"Error processing text file: {str(e)}")
    
    # Calculate processing time
    end_time = datetime.now()
    result['processing_time'] = (end_time - start_time).total_seconds()
    
    return result


def process_doc_file(file_path: str, standardizer) -> Dict:
    """
    Process a legacy DOC file to standardize Bible references.
    Note: This extracts text from DOC files and saves as processed text.
    
    Args:
        file_path: Path to the DOC file
        standardizer: BibleReferenceStandardizer instance
        
    Returns:
        Dictionary with process statistics and status information
    """
    result = {
        'success': False,
        'changes_made': False,
        'paragraphs_processed': 0,
        'paragraphs_changed': 0,
        'backup_path': None,
        'output_path': file_path,
        'error': None,
        'processing_time': 0,
        'warning': 'DOC processing extracts text only. Consider converting to DOCX for full formatting support.'
    }
    
    start_time = datetime.now()
    
    try:
        # Import DOC processing library with fallback
        try:
            import docx2txt
        except ImportError:
            result['error'] = "docx2txt not installed. Install with: pip install docx2txt"
            return result
        
        # Backup the original file
        backup_success, backup_path = standardizer.create_backup(file_path)
        if not backup_success:
            result['error'] = f"Failed to create backup: {backup_path}"
            return result
        result['backup_path'] = backup_path
        
        # Extract text from DOC file
        try:
            text = docx2txt.process(file_path)
        except Exception as e:
            result['error'] = f"Could not extract text from DOC file: {str(e)}"
            return result
        
        if text:
            # Process the extracted text
            new_text, changed = standardizer.process_text(text)
            result['paragraphs_processed'] = 1
            
            if changed:
                result['changes_made'] = True
                result['paragraphs_changed'] = 1
                
                # Save processed text to a new text file
                text_output_path = file_path.replace('.doc', '_processed.txt')
                with open(text_output_path, 'w', encoding='utf-8') as text_file:
                    text_file.write(new_text)
                result['output_path'] = text_output_path
                result['warning'] += f" Processed text saved to: {text_output_path}"
        
        result['success'] = True
        
    except Exception as e:
        result['error'] = str(e)
        logger.error(f"Error processing DOC file: {str(e)}")
    
    # Calculate processing time
    end_time = datetime.now()
    result['processing_time'] = (end_time - start_time).total_seconds()
    
    return result


# Auto-detect file type and process accordingly
def process_any_file(file_path: str, csv_path: Optional[str] = None, 
                    output_path: Optional[str] = None, create_backup: bool = True) -> Dict:
    """
    Automatically detect file type and process accordingly.
    
    Args:
        file_path: Path to the file to process
        csv_path: Path to the CSV file with Bible book abbreviations
        output_path: Path where processed file should be saved
        create_backup: Whether to create a backup
        
    Returns:
        Dictionary with processing results
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    standardizer = BibleReferenceStandardizer(csv_path)
    
    if file_extension == '.docx':
        return standardizer.process_document(file_path, output_path, create_backup)
    elif file_extension == '.html' or file_extension == '.htm':
        return process_html_file(file_path, standardizer)
    elif file_extension == '.pdf':
        return process_pdf_file(file_path, standardizer)
    elif file_extension == '.txt':
        return process_txt_file(file_path, standardizer)
    elif file_extension == '.doc':
        return process_doc_file(file_path, standardizer)
    else:
        return {
            'success': False,
            'error': f"Unsupported file type: {file_extension}. Supported types: .docx, .doc, .html, .pdf, .txt"
        }


# Convenience functions for backward compatibility and ease of use
def process_document(doc_path: str, output_path: Optional[str] = None, 
                    create_backup: bool = True, csv_path: Optional[str] = None) -> Dict:
    """
    Convenience function to process a document with default settings.
    
    Args:
        doc_path: Path to the Word document
        output_path: Path where processed document should be saved
        create_backup: Whether to create a backup
        csv_path: Path to the CSV file with Bible book abbreviations
        
    Returns:
        Dictionary with processing results
    """
    return process_any_file(doc_path, csv_path, output_path, create_backup)


def process_text(text: str, csv_path: Optional[str] = None) -> Tuple[str, bool]:
    """
    Convenience function to process text directly.
    
    Args:
        text: Input text to process
        csv_path: Path to the CSV file with Bible book abbreviations
        
    Returns:
        Tuple of (processed text, whether changes were made)
    """
    standardizer = BibleReferenceStandardizer(csv_path)
    return standardizer.process_text(text)


# Legacy function names for backward compatibility
def process_html(file_path: str, standardizer=None) -> Dict:
    """Legacy function name for HTML processing."""
    if standardizer is None:
        standardizer = BibleReferenceStandardizer()
    return process_html_file(file_path, standardizer)


def process_pdf(file_path: str, standardizer=None) -> Dict:
    """Legacy function name for PDF processing."""
    if standardizer is None:
        standardizer = BibleReferenceStandardizer()
    return process_pdf_file(file_path, standardizer)


def process_txt(file_path: str, standardizer=None) -> Dict:
    """Legacy function name for text processing."""
    if standardizer is None:
        standardizer = BibleReferenceStandardizer()
    return process_txt_file(file_path, standardizer)


def process_doc(file_path: str, standardizer=None) -> Dict:
    """Legacy function name for DOC processing."""
    if standardizer is None:
        standardizer = BibleReferenceStandardizer()
    return process_doc_file(file_path, standardizer)
