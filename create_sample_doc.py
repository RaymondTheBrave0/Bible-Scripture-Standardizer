#!/usr/bin/env python3
"""
Create a sample Word document with various Bible reference formats for testing.
"""

from docx import Document
from docx.shared import Inches

def create_sample_document():
    """Create a sample Word document with various Bible reference formats."""
    
    document = Document()
    
    # Add title
    title = document.add_heading('Sample Document with Bible References', 0)
    
    # Add introduction paragraph
    intro = document.add_paragraph(
        "This document contains various Bible reference formats that need to be standardized. "
        "The following examples show different ways Bible verses are commonly referenced:"
    )
    
    # Add examples
    examples = [
        "In the beginning (gen 1.1), God created the heavens and the earth.",
        "Jesus said in jn 3:16 that God so loved the world.",
        "The famous love chapter is found in 1 cor 13:4-8.",
        "Paul writes in rom 8.28 about all things working together for good.",
        "The Lord's Prayer can be found in matt 6:9-13.",
        "Psalm 23 speaks of the Lord as our shepherd.",
        "In rev 21:4, we read about no more tears.",
        "The Ten Commandments start in ex 20:1.",
        "The Sermon on the Mount begins in Mt 5:3.",
        "Isaiah chapter 53 verse 5 talks about healing.",
        "Consider the lilies as mentioned in luke 12 25-28.",
        "The Great Commission is in Matthew 28:19-20.",
    ]
    
    for example in examples:
        document.add_paragraph(example, style='List Bullet')
    
    # Add a table with more examples
    document.add_heading('Table with Additional References', level=1)
    
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Reference'
    hdr_cells[1].text = 'Topic'
    
    table_data = [
        ('Gen 1:1', 'Creation'),
        ('Ps 23:1', 'The Lord is my shepherd'),
        ('jn 14.6', 'I am the way, truth, and life'),
        ('1 tim 4:12', 'Be an example'),
        ('Phil 4:13', 'I can do all things'),
        ('Jas 1:5', 'Asking for wisdom'),
        ('1 Pet 5:7', 'Cast your cares'),
        ('Heb 11:1', 'Faith definition'),
    ]
    
    for ref, topic in table_data:
        row_cells = table.add_row().cells
        row_cells[0].text = ref
        row_cells[1].text = topic
    
    # Add conclusion
    document.add_heading('Conclusion', level=1)
    conclusion = document.add_paragraph(
        "These references should all be standardized to the format 'Book Chapter:Verse' "
        "such as 'John 3:16' or 'Genesis 1:1'. The standardizer should handle all these "
        "variations automatically."
    )
    
    # Save the document
    filename = 'sample_bible_references.docx'
    document.save(filename)
    print(f"Sample document created: {filename}")
    
    return filename

if __name__ == "__main__":
    create_sample_document()
