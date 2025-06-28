#!/usr/bin/env python3
"""
Create test files in multiple formats to test auto-detection and processing.
"""

import os
from docx import Document

def create_test_files():
    """Create test files in various formats with Bible references."""
    
    # Ensure test directory exists
    os.makedirs('multi_format_tests', exist_ok=True)
    
    # Common test content with Bible references in various formats
    test_content = """
    Bible Study Notes
    
    Today we explored several key passages:
    
    1. Creation Story: gen 1.1 tells us "In the beginning God created..."
    2. God's Love: jn 3:16 is perhaps the most famous verse
    3. The Shepherd Psalm: ps 23 speaks of God's care
    4. Love Chapter: 1 cor 13:4-8 defines love beautifully
    5. All Things Work Together: rom 8.28 gives us hope
    6. The Beatitudes: matt 5:3-12 shows the way to blessing
    7. Great Commission: Matthew 28:19-20 is our calling
    8. Healing Promise: isa chapter 53 verse 5 speaks of healing
    9. Faith Definition: heb 11 1 defines faith
    10. Wisdom Request: jas 1:5 tells us to ask for wisdom
    """
    
    # 1. Create TXT file
    print("Creating text file...")
    with open('multi_format_tests/bible_study.txt', 'w', encoding='utf-8') as f:
        f.write(test_content)
    
    # 2. Create HTML file
    print("Creating HTML file...")
    html_content = f"""<!DOCTYPE html>
<html>
<head>
    <title>Bible Study Notes</title>
    <meta charset="UTF-8">
</head>
<body>
    <h1>Bible Study Notes</h1>
    <p>Today we explored several key passages:</p>
    <ul>
        <li>Creation Story: gen 1.1 tells us "In the beginning God created..."</li>
        <li>God's Love: jn 3:16 is perhaps the most famous verse</li>
        <li>The Shepherd Psalm: ps 23 speaks of God's care</li>
        <li>Love Chapter: 1 cor 13:4-8 defines love beautifully</li>
        <li>All Things Work Together: rom 8.28 gives us hope</li>
        <li>The Beatitudes: matt 5:3-12 shows the way to blessing</li>
        <li>Great Commission: Matthew 28:19-20 is our calling</li>
        <li>Healing Promise: isa chapter 53 verse 5 speaks of healing</li>
        <li>Faith Definition: heb 11 1 defines faith</li>
        <li>Wisdom Request: jas 1:5 tells us to ask for wisdom</li>
    </ul>
    <p>These references demonstrate various formats that need standardization.</p>
</body>
</html>"""
    
    with open('multi_format_tests/bible_study.html', 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    # 3. Create DOCX file
    print("Creating Word document...")
    doc = Document()
    doc.add_heading('Bible Study Notes', 0)
    
    paragraphs = [
        "Today we explored several key passages:",
        "Creation Story: gen 1.1 tells us 'In the beginning God created...'",
        "God's Love: jn 3:16 is perhaps the most famous verse",
        "The Shepherd Psalm: ps 23 speaks of God's care",
        "Love Chapter: 1 cor 13:4-8 defines love beautifully",
        "All Things Work Together: rom 8.28 gives us hope",
        "The Beatitudes: matt 5:3-12 shows the way to blessing",
        "Great Commission: Matthew 28:19-20 is our calling",
        "Healing Promise: isa chapter 53 verse 5 speaks of healing",
        "Faith Definition: heb 11 1 defines faith",
        "Wisdom Request: jas 1:5 tells us to ask for wisdom"
    ]
    
    for para in paragraphs:
        doc.add_paragraph(para)
    
    doc.save('multi_format_tests/bible_study.docx')
    
    # 4. Create a simple PDF-like content file (we'll treat as text for now)
    print("Creating additional text file for PDF simulation...")
    with open('multi_format_tests/sermon_notes.txt', 'w', encoding='utf-8') as f:
        f.write("""
        Sunday Sermon: Faith and Trust
        
        Key verses discussed:
        - Trust: prov 3:5-6 about trusting in the Lord
        - Faith: mark 11.24 about believing when you pray
        - Hope: jer 29:11 about God's plans for us
        - Peace: phil 4:6-7 about prayer and peace
        - Strength: isa 40.31 about waiting on the Lord
        
        These passages show different formatting styles.
        """)
    
    print("\nTest files created in 'multi_format_tests/' directory:")
    print("  - bible_study.txt")
    print("  - bible_study.html") 
    print("  - bible_study.docx")
    print("  - sermon_notes.txt")

if __name__ == "__main__":
    create_test_files()
